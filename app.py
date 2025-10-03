# -*- coding: utf-8 -*-
import os
import io
import hashlib
from datetime import datetime

import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# ========================
#  安全取得 PEPPER（本機免 secrets 也能跑）
# ========================
def sha256_hex(s: str) -> str:
    return hashlib.sha256(s.encode("utf-8")).hexdigest()

def get_pepper():
    val = os.getenv("PEPPER")
    if val:
        return val
    try:
        if hasattr(st, "secrets") and "PEPPER" in st.secrets:
            return st.secrets["PEPPER"]
    except Exception:
        pass
    return "oQ3s7_Pepp3r!2025#"  # fallback，正式請改用環境變數或 secrets

PEPPER = get_pepper()

# ========================
#  使用者與雜湊（只存雜湊，不存明文）
#  Hash 公式：SHA256( PEPPER + password + SALT )
# ========================
USERS = {
    "Charles": {"name": "Charles", "salt": "HxEFDUaRey", "pw_hash": "d1b6611d51176cb4d445fc8d98debcf7d09afba09c84681ec4e0b8560438a9d4"},
    "Hsiang":  {"name": "Hsiang",  "salt": "oNeiHDVCSY", "pw_hash": "a70d3210d4960b5d3a502a58d14d37c6a5eb66c036c5301b94259fa3bd6ad8ba"},
    "Sandy":   {"name": "Sandy",   "salt": "s1pXUeWFmY", "pw_hash": "36aff05ac4953b773dee06df6dbc2850f3b54c0748f10d36ab4b7e9cb73062e2"},
    "Min":     {"name": "Min",     "salt": "5OaOjhx3Ty", "pw_hash": "2b8997a0ff9e3d99d37fdd0de9b27b9c2a959ed18f8b68a5e774bb3d44999583"},
    "May":     {"name": "May",     "salt": "oKNt3n4onv", "pw_hash": "a3d5445507e88eceb576da31f5750fa0fd826bd354396ce2414a6f54e7648898"},
    "Ping":    {"name": "Ping",    "salt": "KXCbcUEJaf", "pw_hash": "1165ba04fa78320f59a286ade424fbfdfd47975269bbbba70697a3e2f685488a"},
    "Denny":   {"name": "Denny",   "salt": "81glcT1w6J", "pw_hash": "5e0333bd0d02df45dd9d5ae7deac35f8033c9db0defd6cd4551f2397b082b98f"},
}

def check_login(username: str, password: str) -> bool:
    u = USERS.get(username)
    if not u:
        return False
    calc = sha256_hex(PEPPER + password + u["salt"])
    return calc == u["pw_hash"]

# 「宣達人」對應工號（用來在名單中停用該人）
SPOKESMAN_TO_ID = {"陳淑敏": "B00011", "陳玫曄": "B00013", "郭秀坪": "B00039"}

# ========================
#  名單（已移除 B00025/B00017/B00040/B00054）
# ========================
ALL_PERSONNEL = [
    ("B00011", "陳淑敏"), ("B00013", "陳玫曄"), ("B00015", "羅思如"),
    ("B00018", "黃碧苓"), ("B00019", "劉文斌"),
    ("B00022", "江淑敏"), ("B00034", "黃慧珊"),
    ("B00038", "吳沛璇"), ("B00039", "郭秀坪"),
    ("B00052", "靳椀婷"), ("B00065", "沈玉婷"),
    ("B00074", "江憶嬋"), ("B00075", "張珮珊"), ("B00082", "林鈺蓉"),
    ("B00085", "彭雅芬"), ("B00088", "劉諭潔"),("B00098", "徐淑惠"),
    ("F00001", "黛安娜"), ("F00002", "雪莉"), ("F00003", "喬得琳"),
    ("F00004", "妮琪"), ("F00005", "艾莉安"), ("F00006", "裘蒂"),
    ("F00007", "羅莎琳"), ("F00008", "安喬"), ("F00009", "玫貞"),
    ("F00010", "媚拉"), ("F00011", "姍米"), ("F00012", "艾達"),
    ("F00013", "喬比"), ("F00014", "娜琳"), ("F00015", "芙羅娜"),
    ("F00016", "愛拉"), ("F00017", "裘娜琳"), ("F00018", "雪拉"),
    ("F00019", "克利絲"), ("F00020", "洛娜"), ("F00021", "茱莉"),
    ("F00022", "希拉"), ("F00023", "瑪琳"), ("F00024", "潔絲汀"),
    ("F00025", "瑪德琳"), ("F00026", "艾琳"), ("F00027", "瑪莉"),
    ("F00028", "丹妮"), ("F00029", "瑪芮"), ("F00030", "羅貝玲"),
    ("F00031", "凱莎"),
]

# ========================
#  Word 寫入工具
# ========================
def set_paragraph_font(paragraph, name="標楷體", size=12):
    for r in paragraph.runs:
        r.font.name = name
        r._element.rPr.rFonts.set(qn("w:eastAsia"), name)
        r.font.size = Pt(size)

def set_doc_field_by_label(doc, label_keywords, value, font_name="標楷體", font_size=12):
    """尋找含標籤文字的儲存格，將右側（或下一列同欄）覆寫為 value。"""
    def hit(text):
        t = (text or "").replace("\n", " ").strip()
        return any(k in t for k in label_keywords)

    for tbl in doc.tables:
        rows = tbl.rows
        for ri, row in enumerate(rows):
            cells = row.cells
            for ci, cell in enumerate(cells):
                if hit(cell.text):
                    tgt = cells[ci + 1] if ci + 1 < len(cells) else None
                    if tgt is None and ri + 1 < len(rows):
                        tgt = rows[ri + 1].cells[ci]
                    if tgt is not None:
                        tgt.text = ""
                        p = tgt.paragraphs[0]
                        r = p.add_run(value)
                        r.font.name = font_name
                        r._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
                        r.font.size = Pt(font_size)
                        return True
    return False

def build_word(template_bytes: bytes, attendees, meta):
    doc = Document(io.BytesIO(template_bytes))

    # 寫「資訊表格」
    set_doc_field_by_label(doc, ["地點", "Location"], meta.get("location", ""))
    set_doc_field_by_label(doc, ["日期", "Date"],     meta.get("date", ""))
    set_doc_field_by_label(doc, ["時間", "Time"],     meta.get("time", ""))
    set_doc_field_by_label(doc, ["宣達人", "spokesman", "Spokesman"], meta.get("spokesman", ""))

    # 出席名單：預設取 tables[1]
    try:
        target_table = doc.tables[1]
    except IndexError:
        raise RuntimeError("範本格式異常：找不到出席名單表（tables[1]）。請確認範本。")

    # 清空舊資料：保留表頭列
    for _ in range(len(target_table.rows) - 1):
        target_table._tbl.remove(target_table.rows[1]._tr)

    # 逐列寫入
    for emp_id, name in attendees:
        row = target_table.add_row()
        c_id, c_name, c_sign, c_date = row.cells
        c_id.text = emp_id
        c_name.text = name
        c_sign.text = ""
        c_date.text = ""

        for p in c_id.paragraphs:
            set_paragraph_font(p, name="Arial", size=12)
            for r in p.runs:
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        for p in c_name.paragraphs:
            set_paragraph_font(p, name="標楷體", size=12)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# ========================
#  設定
# ========================
st.set_page_config(page_title="生產線出席勾選（雲端版）", layout="wide")
TEMPLATE_PATH = "Template/生產線每日宣達事項_PD_範本.docx"  # << 你把範本放這裡

# --- Session ---
if "auth" not in st.session_state:
    st.session_state.auth = False
if "user" not in st.session_state:
    st.session_state.user = None
if "checked" not in st.session_state:
    st.session_state.checked = set(eid for eid, _ in ALL_PERSONNEL)  # 預設全選

# --- Login（帳號下拉 + 密碼） ---
with st.sidebar:
    st.header("登入")
    if not st.session_state.auth:
        usernames = list(USERS.keys())
        sel_user = st.selectbox("帳號（姓名）", usernames, index=0)
        password = st.text_input("密碼（你的工號）", type="password", key="p")
        if st.button("登入", use_container_width=True):
            if check_login(sel_user, password.strip()):
                st.session_state.auth = True
                st.session_state.user = sel_user
                st.success(f"歡迎 {USERS[sel_user]['name']}！")
            else:
                st.error("帳號或密碼錯誤")
    else:
        st.success(f"已登入：{st.session_state.user}")
        if st.button("登出", use_container_width=True):
            st.session_state.auth = False
            st.session_state.user = None
            st.session_state.checked = set(eid for eid, _ in ALL_PERSONNEL)
            st.rerun()

st.title("每日宣達出席紀錄表（雲端版）")
if not st.session_state.auth:
    st.info("請先於左側登入。")
    st.stop()

# --- 宣達資訊 ---
with st.expander("① 宣達資訊（必填）", expanded=True):
    c1, c2, c3, c4 = st.columns([1.2, 1, 1, 1])
    location = c1.selectbox("地點 / Location", ["工四廠-C1", "工四廠", "工四廠-A/B"], index=1)
    date_str = c2.text_input("日期 / Date（YYYY/MM/DD）", value=datetime.now().strftime("%Y/%m/%d"))
    time_str = c3.text_input("時間 / Time", value="08:00 ~ 08:15")
    spokesman = c4.selectbox("宣達人 / Spokesman", ["陳淑敏", "陳玫曄", "郭秀坪"], index=0)

# 取得「宣達人」的工號，並確保不可被勾選
spokesman_id = SPOKESMAN_TO_ID.get(spokesman)
if spokesman_id in st.session_state.checked:
    st.session_state.checked.discard(spokesman_id)

# --- 範本（改為固定路徑） ---
with st.expander("② 範本路徑（預設即可）", expanded=True):
    cpath1, cpath2 = st.columns([3, 1])
    tpath = cpath1.text_input("範本檔路徑（.docx）", value=TEMPLATE_PATH)
    cpath2.button("重新整理", help="若你剛上傳/更新範本檔到專案，按一下重新整理頁面", on_click=lambda: st.rerun())
    # 存回變數（允許你改路徑）
    TEMPLATE_PATH = tpath

    if not os.path.exists(TEMPLATE_PATH):
        st.error(f"找不到範本：{TEMPLATE_PATH}\n請確認檔案已放入專案對應路徑。")
    else:
        st.success(f"已找到範本：{TEMPLATE_PATH}")

# --- 名單（單頁可視，宣達人自動停用） ---
with st.expander("③ 勾選出席人員", expanded=True):
    ctop1, ctop2, ctop3, ctop4, ctop5 = st.columns([1.6, 1, 1, 1, 1])
    query = ctop1.text_input("搜尋（可輸入工號或姓名片段）", "")
    show_b = ctop2.toggle("只看本籍 (B)", value=False)
    show_f = ctop3.toggle("只看外籍 (F)", value=False)
    num_cols = ctop4.slider("每列欄數", min_value=2, max_value=8, value=6, help="調大可讓更多人一頁顯示")
    precheck_all = ctop5.checkbox("將目前可見 → 全部勾選", value=False)

    filtered = []
    for emp_id, name in ALL_PERSONNEL:
        if query and (query not in emp_id and query not in name):
            continue
        if show_b and not emp_id.startswith("B"):
            continue
        if show_f and not emp_id.startswith("F"):
            continue
        filtered.append((emp_id, name))

    b1, b2, b3 = st.columns(3)
    if b1.button("目前可見 → 全部勾選"):
        for emp_id, _ in filtered:
            if emp_id != spokesman_id:
                st.session_state.checked.add(emp_id)
    if b2.button("目前可見 → 全部清除"):
        for emp_id, _ in filtered:
            st.session_state.checked.discard(emp_id)
    if b3.button("全部清除（含不可見）"):
        st.session_state.checked.clear()

    if precheck_all:
        for emp_id, _ in filtered:
            if emp_id != spokesman_id:
                st.session_state.checked.add(emp_id)

    st.caption(f"提示：宣達人「{spokesman}」已自動排除，無法選為出席人員。")

    rows = (len(filtered) + num_cols - 1) // num_cols
    idx = 0
    for _ in range(rows):
        cols = st.columns(num_cols)
        for col in cols:
            if idx >= len(filtered):
                break
            emp_id, name = filtered[idx]
            label = f"{emp_id}  {name}"
            if emp_id == spokesman_id:
                col.checkbox(label + "（宣達人，自動排除）", value=False, disabled=True, key=f"chk_{emp_id}")
                st.session_state.checked.discard(emp_id)
            else:
                checked = emp_id in st.session_state.checked
                if col.checkbox(label, value=checked, key=f"chk_{emp_id}"):
                    st.session_state.checked.add(emp_id)
                else:
                    st.session_state.checked.discard(emp_id)
            idx += 1

# --- 產出 Word ---
st.divider()
cL, cR = st.columns([3, 1])
with cL:
    st.caption("完成後可下載 Word 檔。")
gen_btn = cR.button("④ 產生出席記錄", use_container_width=True)

if gen_btn:
    if not os.path.exists(TEMPLATE_PATH):
        st.error("請先確認範本路徑正確，檔案存在。")
    else:
        if not date_str.strip():
            st.error("請輸入日期（YYYY/MM/DD）")
            st.stop()
        if "~" not in time_str:
            st.error("時間格式請用區間，例如：08:00 ~ 08:15")
            st.stop()

        selected = [(emp_id, name) for emp_id, name in ALL_PERSONNEL if emp_id in st.session_state.checked]
        if not selected:
            st.error("請至少勾選一位人員")
            st.stop()

        meta = {
            "location": location.strip(),
            "date": date_str.strip(),
            "time": time_str.strip(),
            "spokesman": spokesman.strip(),
        }

        try:
            with open(TEMPLATE_PATH, "rb") as f:
                tmpl_bytes = f.read()
            out_io = build_word(tmpl_bytes, selected, meta)
            ts = datetime.now().strftime("%Y%m%d-%H%M")
            out_name = f"生產線每日宣達事項_出席記錄_{ts}.docx"
            st.success("已產生完成！")
            st.download_button(
                "下載 Word 檔",
                data=out_io.getvalue(),
                file_name=out_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:
            st.exception(e)

